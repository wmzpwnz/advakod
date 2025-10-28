"""
Сервис для работы с документами
Загрузка, парсинг и индексация документов для RAG системы
"""

import logging
import os
import asyncio
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
import uuid
from datetime import datetime
import tempfile
import requests
from urllib.parse import urlparse
import mimetypes

# Импорты для работы с разными форматами файлов
import pypdf
import pdfplumber
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .embeddings_service import embeddings_service
from .vector_store_service import vector_store_service
from .document_validator import document_validator
from .ai_document_validator import ai_document_validator
from .hybrid_document_validator import hybrid_document_validator
from .document_versioning import document_versioning_service
from .simple_expert_rag import simple_expert_rag
from .pdf_ocr_service import pdf_ocr_service

logger = logging.getLogger(__name__)

class DocumentService:
    """Сервис для работы с документами"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # Размер чанка в символах
            chunk_overlap=200,  # Перекрытие между чанками
            length_function=len,
            separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
        )
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
        self.validation_method = "none"  # hybrid, ai, rules, none
        
    def get_status(self) -> Dict[str, Any]:
        """Возвращает статус сервиса"""
        return {
            "supported_extensions": list(self.supported_extensions),
            "chunk_size": self.text_splitter._chunk_size,
            "chunk_overlap": self.text_splitter._chunk_overlap,
            "validation_method": self.validation_method,
            "hybrid_stats": hybrid_document_validator.get_validation_stats()
        }
    
    def set_validation_method(self, method: str):
        """Переключает метод валидации"""
        if method not in ["hybrid", "ai", "rules", "none"]:
            raise ValueError("Метод валидации должен быть: hybrid, ai, rules или none")
        
        self.validation_method = method
        logger.info(f"Метод валидации изменен на: {method}")
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Вычисляет хеш файла для проверки дубликатов"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _get_pdf_page_count(self, file_path: str) -> int:
        """Получает количество страниц в PDF файле"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            logger.warning(f"Не удалось подсчитать страницы PDF {file_path}: {e}")
            return 0
    
    def _get_docx_page_count(self, file_path: str) -> int:
        """Получает приблизительное количество страниц в DOCX файле"""
        try:
            doc = Document(file_path)
            # Подсчитываем количество параграфов и делим на примерное количество параграфов на страницу
            paragraphs = len(doc.paragraphs)
            # Примерно 20-30 параграфов на страницу для юридических документов
            estimated_pages = max(1, paragraphs // 25)
            return estimated_pages
        except Exception as e:
            logger.warning(f"Не удалось подсчитать страницы DOCX {file_path}: {e}")
            return 0
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Извлекает текст из PDF файла с использованием разных методов"""
        try:
            # Метод 1: pdfplumber (более мощный)
            text = self._extract_text_with_pdfplumber(file_path)
            if text and len(text.strip()) > 100:
                logger.info(f"pdfplumber извлечение успешно: {len(text)} символов")
                return text.strip()
            
            # Метод 2: pypdf (стандартный)
            text = self._extract_text_with_pypdf(file_path)
            if text and len(text.strip()) > 100:
                logger.info(f"pypdf извлечение успешно: {len(text)} символов")
                return text.strip()
            
            # Метод 3: OCR (если другие не сработали)
            logger.info("Стандартные методы не дали результата, пробуем OCR")
            ocr_text = pdf_ocr_service.extract_text_from_pdf(file_path)
            
            if ocr_text and len(ocr_text.strip()) > 50:
                logger.info(f"OCR извлечение успешно: {len(ocr_text)} символов")
                return ocr_text.strip()
            else:
                logger.warning("Все методы извлечения текста не дали результата")
                return text.strip() if text else ""
                
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из PDF {file_path}: {e}")
            return ""
    
    def _extract_text_with_pdfplumber(self, file_path: str) -> str:
        """Извлекает текст из PDF с помощью pdfplumber"""
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Ошибка извлечения текста со страницы {page_num} (pdfplumber): {e}")
                        continue
            return '\n\n'.join(text_parts) if text_parts else ""
        except Exception as e:
            logger.error(f"Ошибка pdfplumber извлечения: {e}")
            return ""
    
    def _extract_text_with_pypdf(self, file_path: str) -> str:
        """Извлекает текст из PDF с помощью pypdf"""
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Ошибка извлечения текста со страницы {page_num} (pypdf): {e}")
                        continue
            return '\n\n'.join(text_parts) if text_parts else ""
        except Exception as e:
            logger.error(f"Ошибка pypdf извлечения: {e}")
            return ""
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            logger.info(f"🔍 Начинаем извлечение текста из DOCX: {file_path}")
            doc = Document(file_path)
            text = ""
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Пропускаем пустые параграфы
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            logger.info(f"📊 Извлечено {paragraph_count} параграфов, {len(text)} символов")
            result = text.strip()
            logger.info(f"✅ DOCX извлечение завершено: {len(result)} символов")
            return result
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения текста из DOCX {file_path}: {e}")
            logger.error(f"Тип ошибки: {type(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return ""
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Извлекает текст из TXT файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Пробуем другие кодировки
            for encoding in ['cp1251', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except:
                    continue
            logger.error(f"Не удалось определить кодировку файла {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Ошибка чтения TXT файла {file_path}: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Извлекает текст из файла в зависимости от его типа"""
        if not os.path.exists(file_path):
            logger.error(f"Файл не найден: {file_path}")
            return ""
        
        file_extension = Path(file_path).suffix.lower()
        logger.info(f"📄 Обрабатываем файл: {file_path} (расширение: {file_extension})")
        
        if file_extension not in self.supported_extensions:
            logger.error(f"Неподдерживаемый формат файла: {file_extension}")
            logger.error(f"Поддерживаемые форматы: {self.supported_extensions}")
            return ""
        
        logger.info(f"📄 Извлекаем текст из файла: {file_path}")
        
        if file_extension == '.pdf':
            result = self._extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            result = self._extract_text_from_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            result = self._extract_text_from_txt(file_path)
        else:
            logger.error(f"Обработчик для {file_extension} не реализован")
            return ""
        
        logger.info(f"📊 Результат извлечения: {len(result)} символов")
        return result
    
    def split_text_into_chunks(self, text: str) -> List[str]:
        """Разбивает текст на чанки для индексации"""
        if not text.strip():
            return []
        
        try:
            chunks = self.text_splitter.split_text(text)
            logger.info(f"📝 Текст разбит на {len(chunks)} чанков")
            return chunks
        except Exception as e:
            logger.error(f"Ошибка разбивки текста на чанки: {e}")
            return [text]  # Возвращаем исходный текст как один чанк
    
    async def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Обрабатывает файл и добавляет его в векторную базу данных"""
        logger.info(f"🚀 Начинаем обработку файла: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"❌ Файл не найден: {file_path}")
            return {"success": False, "error": f"Файл не найден: {file_path}"}
        
        try:
            # Извлекаем текст
            logger.info(f"📄 Извлекаем текст из файла...")
            text = self.extract_text_from_file(file_path)
            logger.info(f"📊 Извлеченный текст: {len(text)} символов")
            
            if not text:
                logger.error("❌ Не удалось извлечь текст из файла")
                return {"success": False, "error": "Не удалось извлечь текст из файла"}
            
            # Валидируем документ на соответствие юридической тематике
            file_info = Path(file_path)
            logger.info(f"🔍 Валидируем документ методом: {self.validation_method}")
            
            if self.validation_method == "none":
                # Валидация отключена - принимаем все документы
                validation_result = {
                    "is_valid": True,
                    "document_type": "unknown",
                    "confidence": 1.0,
                    "reason": "Валидация отключена",
                    "legal_score": 1.0,
                    "invalid_score": 0.0
                }
            elif self.validation_method == "hybrid":
                # Используем гибридную валидацию
                validation_result = await hybrid_document_validator.validate_document(text, file_info.name)
            elif self.validation_method == "ai":
                # Используем AI валидацию
                validation_result = await ai_document_validator.validate_document(text, file_info.name)
            else:  # rules
                # Используем обычную валидацию
                validation_result = document_validator.validate_document(text, file_info.name)
            
            logger.info(f"📊 Результат валидации: {validation_result}")
            
            if not validation_result["is_valid"]:
                logger.warning(f"⚠️ Документ не прошел валидацию: {validation_result['reason']}")
                return {
                    "success": False, 
                    "error": f"Документ не соответствует юридической тематике: {validation_result['reason']}",
                    "validation_details": validation_result,
                    "suggestions": validation_result.get("suggestions", [])
                }
            
            # Разбиваем на чанки
            logger.info(f"✂️ Разбиваем текст на чанки...")
            chunks = self.split_text_into_chunks(text)
            logger.info(f"📊 Создано чанков: {len(chunks)}")
            
            if not chunks:
                logger.error("❌ Не удалось разбить текст на чанки")
                return {"success": False, "error": "Не удалось разбить текст на чанки"}
            
            # Подготавливаем метаданные
            file_info = Path(file_path)
            file_hash = self._calculate_file_hash(file_path)
            
            # Подсчитываем количество страниц
            pages_count = 0
            if file_info.suffix.lower() == '.pdf':
                pages_count = self._get_pdf_page_count(file_path)
            elif file_info.suffix.lower() == '.docx':
                pages_count = self._get_docx_page_count(file_path)
            
            # Создаем версию документа
            document_id = f"doc_{file_hash[:8]}"
            version = document_versioning_service.create_document_version(
                document_id=document_id,
                text=text,
                filename=file_info.name,
                file_hash=file_hash,
                metadata={"original_filename": file_info.name}
            )
            
            # Упрощенные метаданные для ChromaDB (только простые типы)
            base_metadata = {
                "filename": file_info.name,
                "file_extension": file_info.suffix.lower(),
                "file_size": int(os.path.getsize(file_path)),
                "file_hash": file_hash,
                "processed_at": datetime.now().isoformat(),
                "chunks_count": int(len(chunks)),
                "total_length": int(len(text)),
                "document_type": str(validation_result["document_type"]),
                "validation_confidence": float(validation_result["confidence"]),
                "legal_score": float(validation_result.get("legal_score", 0.0)),
                "is_validated": True,
                "document_id": document_id,
                "version": str(version.version),
                "status": str(version.status),
                "is_draft": version.status == "draft",
                "pages": int(pages_count) if pages_count > 0 else None
            }
            
            # Фильтруем None значения из базовых метаданных
            base_metadata = {k: v for k, v in base_metadata.items() if v is not None}
            
            if metadata:
                # Фильтруем None значения из дополнительных метаданных
                filtered_metadata = {k: v for k, v in metadata.items() if v is not None}
                base_metadata.update(filtered_metadata)
            
            # Инициализируем векторную базу данных если не готова
            if not vector_store_service.is_ready():
                logger.info("Инициализируем векторную базу данных...")
                vector_store_service.initialize()
                logger.info(f"Vector store ready after init: {vector_store_service.is_ready()}")
            
            # Добавляем чанки в векторную базу данных
            added_count = 0
            logger.info(f"Начинаем добавление {len(chunks)} чанков в векторную БД")
            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "chunk_index": int(i),
                    "chunk_length": int(len(chunk)),
                    "is_chunk": True
                })
                
                # Фильтруем None значения из метаданных
                chunk_metadata = {k: v for k, v in chunk_metadata.items() if v is not None}
                
                chunk_id = f"{file_hash}_{i}"
                logger.info(f"Добавляем чанк {i+1}/{len(chunks)}: {chunk_id}")
                
                # Проверяем готовность векторной БД перед каждым добавлением
                if not vector_store_service.is_ready():
                    logger.warning(f"Vector store не готов для чанка {i+1}, инициализируем...")
                    vector_store_service.initialize()
                
                success = vector_store_service.add_document(
                    content=chunk,
                    metadata=chunk_metadata,
                    document_id=chunk_id
                )
                logger.info(f"Результат добавления чанка {i+1}: {success}")
                
                if success:
                    added_count += 1
            
            logger.info(f"✅ Файл обработан: {file_path} ({added_count}/{len(chunks)} чанков)")
            
            # Добавляем документ в simple_expert_rag
            try:
                from .simple_expert_rag import LegalMetadata
                
                # Создаем метаданные для simple_expert_rag
                legal_metadata = LegalMetadata(
                    source=file_info.name,
                    ingested_at=datetime.now()
                )
                
                # Добавляем документ в simple_expert_rag
                simple_rag_result = await simple_expert_rag.add_document_with_text(
                    text_content=text,
                    metadata=legal_metadata
                )
                
                if simple_rag_result.get('success'):
                    logger.info(f"✅ Документ добавлен в simple_expert_rag: {simple_rag_result.get('chunks_created', 0)} чанков")
                else:
                    logger.warning(f"⚠️ Ошибка добавления в simple_expert_rag: {simple_rag_result.get('error', 'Unknown error')}")
                    
            except Exception as e:
                logger.error(f"❌ Ошибка интеграции с simple_expert_rag: {e}")
            
            return {
                "success": True,
                "file_path": file_path,
                "chunks_added": added_count,
                "total_chunks": len(chunks),
                "file_hash": file_hash,
                "text_length": len(text),
                "document_id": document_id
            }
            
        except Exception as e:
            logger.error(f"❌ Ошибка обработки файла {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    async def process_directory(self, directory_path: str, recursive: bool = True) -> Dict[str, Any]:
        """Обрабатывает все поддерживаемые файлы в директории"""
        if not os.path.exists(directory_path):
            return {"success": False, "error": f"Директория не найдена: {directory_path}"}
        
        files_found = []
        
        # Ищем файлы
        if recursive:
            for root, dirs, files in os.walk(directory_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    if Path(file_path).suffix.lower() in self.supported_extensions:
                        files_found.append(file_path)
        else:
            for file in os.listdir(directory_path):
                file_path = os.path.join(directory_path, file)
                if os.path.isfile(file_path) and Path(file_path).suffix.lower() in self.supported_extensions:
                    files_found.append(file_path)
        
        if not files_found:
            return {"success": False, "error": "Поддерживаемые файлы не найдены"}
        
        logger.info(f"📁 Найдено файлов для обработки: {len(files_found)}")
        
        # Обрабатываем файлы
        results = []
        for file_path in files_found:
            result = await self.process_file(file_path)
            results.append(result)
        
        successful = [r for r in results if r.get("success", False)]
        failed = [r for r in results if not r.get("success", False)]
        
        return {
            "success": True,
            "directory_path": directory_path,
            "files_found": len(files_found),
            "files_processed": len(successful),
            "files_failed": len(failed),
            "results": results
        }
    
    async def add_text_document(self, 
                              title: str, 
                              content: str, 
                              metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Добавляет текстовый документ напрямую"""
        if not content.strip():
            return {"success": False, "error": "Пустой контент"}
        
        try:
            # Валидируем документ на соответствие юридической тематике
            if self.validation_method == "none":
                # Валидация отключена - принимаем все документы
                validation_result = {
                    "is_valid": True,
                    "document_type": "unknown",
                    "confidence": 1.0,
                    "reason": "Валидация отключена",
                    "legal_score": 1.0,
                    "invalid_score": 0.0
                }
            elif self.validation_method == "hybrid":
                # Используем гибридную валидацию
                validation_result = await hybrid_document_validator.validate_document(content, title)
            elif self.validation_method == "ai":
                # Используем AI валидацию
                validation_result = await ai_document_validator.validate_document(content, title)
            else:  # rules
                # Используем обычную валидацию
                validation_result = document_validator.validate_document(content, title)
            
            if not validation_result["is_valid"]:
                return {
                    "success": False, 
                    "error": f"Документ не соответствует юридической тематике: {validation_result['reason']}",
                    "validation_details": validation_result,
                    "suggestions": validation_result.get("suggestions", [])
                }
            
            # Разбиваем на чанки
            chunks = self.split_text_into_chunks(content)
            if not chunks:
                return {"success": False, "error": "Не удалось разбить текст на чанки"}
            
            # Подготавливаем метаданные
            doc_id = str(uuid.uuid4())
            content_hash = hashlib.sha256(content.encode()).hexdigest()
            
            base_metadata = {
                "title": title,
                "document_id": doc_id,
                "content_hash": content_hash,
                "added_at": datetime.now().isoformat(),
                "chunks_count": len(chunks),
                "total_length": len(content),
                "source": "direct_input",
                # Добавляем информацию о валидации
                "document_type": validation_result["document_type"],
                "validation_confidence": validation_result["confidence"],
                "legal_score": validation_result.get("legal_score", 0.0),
                "is_validated": True
            }
            
            if metadata:
                base_metadata.update(metadata)
            
            # Инициализируем векторную базу данных если не готова
            if not vector_store_service.is_ready():
                logger.info("Инициализируем векторную базу данных...")
                vector_store_service.initialize()
            
            # Добавляем чанки в векторную базу данных
            added_count = 0
            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "chunk_index": i,
                    "chunk_length": len(chunk),
                    "is_chunk": True
                })
                
                chunk_id = f"{doc_id}_{i}"
                success = vector_store_service.add_document(
                    content=chunk,
                    metadata=chunk_metadata,
                    document_id=chunk_id
                )
                
                if success:
                    added_count += 1
            
            logger.info(f"✅ Текстовый документ добавлен: {title} ({added_count}/{len(chunks)} чанков)")
            
            return {
                "success": True,
                "title": title,
                "document_id": doc_id,
                "chunks_added": added_count,
                "total_chunks": len(chunks),
                "content_hash": content_hash
            }
            
        except Exception as e:
            logger.error(f"❌ Ошибка добавления текстового документа: {e}")
            return {"success": False, "error": str(e)}

    async def process_url(self, url: str) -> Dict[str, Any]:
        """Обработка документа по URL"""
        try:
            logger.info(f"🌐 Начинаем обработку URL: {url}")
            
            # Валидация URL
            parsed_url = urlparse(url)
            if not parsed_url.scheme or not parsed_url.netloc:
                return {"success": False, "error": "Некорректный URL"}
            
            # Определяем тип контента по URL
            content_type = self._get_content_type_from_url(url)
            logger.info(f"📄 Определен тип контента: {content_type}")
            
            # Загружаем контент
            if content_type == "html":
                text = await self._extract_text_from_html_url(url)
            else:
                text = await self._extract_text_from_file_url(url)
            
            if not text:
                return {"success": False, "error": "Не удалось извлечь текст из URL"}
            
            logger.info(f"📝 Извлечено {len(text)} символов из URL")
            
            # Создаем временный файл для обработки
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as tmp_file:
                tmp_file.write(text)
                tmp_file_path = tmp_file.name
            
            try:
                # Обрабатываем как текстовый документ
                result = await self.add_text_document(
                    title=f"url_document_{uuid.uuid4().hex[:8]}.txt",
                    content=text,
                    metadata={
                        "source_type": "url",
                        "source_url": url,
                        "content_type": content_type,
                        "downloaded_at": datetime.now().isoformat()
                    }
                )
                
                if result.get('success'):
                    result['source_url'] = url
                    result['content_type'] = content_type
                
                return result
                
            finally:
                # Удаляем временный файл
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                    
        except Exception as e:
            logger.error(f"❌ Ошибка обработки URL {url}: {e}")
            return {"success": False, "error": str(e)}
    
    def _get_content_type_from_url(self, url: str) -> str:
        """Определяет тип контента по URL"""
        parsed_url = urlparse(url)
        path = parsed_url.path.lower()
        
        # Проверяем расширение файла
        if path.endswith('.pdf'):
            return 'pdf'
        elif path.endswith(('.doc', '.docx')):
            return 'docx'
        elif path.endswith(('.txt', '.text')):
            return 'txt'
        elif path.endswith(('.md', '.markdown')):
            return 'md'
        else:
            # По умолчанию считаем HTML
            return 'html'
    
    async def _extract_text_from_html_url(self, url: str) -> str:
        """Извлекает текст из HTML страницы"""
        try:
            logger.info(f"🌐 Загружаем HTML контент: {url}")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Устанавливаем правильную кодировку
            response.encoding = response.apparent_encoding or 'utf-8'
            
            # Улучшенное извлечение текста из HTML
            import re
            html_content = response.text
            
            # Удаляем скрипты и стили
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            
            # Удаляем HTML теги
            clean_text = re.sub(r'<[^>]+>', '', html_content)
            
            # Декодируем HTML entities
            import html
            clean_text = html.unescape(clean_text)
            
            # Удаляем лишние пробелы и переносы строк
            clean_text = re.sub(r'\s+', ' ', clean_text)
            clean_text = clean_text.strip()
            
            logger.info(f"📄 Извлечено {len(clean_text)} символов из HTML")
            return clean_text
            
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения текста из HTML {url}: {e}")
            return ""
    
    async def _extract_text_from_file_url(self, url: str) -> str:
        """Извлекает текст из файла по URL"""
        try:
            logger.info(f"📁 Загружаем файл: {url}")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Определяем тип файла по заголовкам
            content_type = response.headers.get('content-type', '').lower()
            
            if 'pdf' in content_type or url.lower().endswith('.pdf'):
                return await self._extract_text_from_pdf_bytes(response.content)
            elif 'msword' in content_type or 'officedocument' in content_type or url.lower().endswith(('.doc', '.docx')):
                return await self._extract_text_from_docx_bytes(response.content)
            else:
                # Пытаемся извлечь как текст
                try:
                    text = response.text
                    logger.info(f"📄 Извлечено {len(text)} символов как текст")
                    return text
                except:
                    return ""
                    
        except Exception as e:
            logger.error(f"❌ Ошибка загрузки файла {url}: {e}")
            return ""
    
    async def _extract_text_from_pdf_bytes(self, pdf_bytes: bytes) -> str:
        """Извлекает текст из PDF байтов"""
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_file.write(pdf_bytes)
                tmp_file_path = tmp_file.name
            
            try:
                return await self._extract_text_from_pdf(tmp_file_path)
            finally:
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения текста из PDF байтов: {e}")
            return ""
    
    async def _extract_text_from_docx_bytes(self, docx_bytes: bytes) -> str:
        """Извлекает текст из DOCX байтов"""
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(docx_bytes)
                tmp_file_path = tmp_file.name
            
            try:
                return await self._extract_text_from_docx(tmp_file_path)
            finally:
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения текста из DOCX байтов: {e}")
            return ""

# Глобальный экземпляр сервиса
document_service = DocumentService()
