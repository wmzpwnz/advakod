import logging
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
import uuid
import json
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import openpyxl
from docx import Document
from docx.shared import Inches

logger = logging.getLogger(__name__)

class ExportFormat(Enum):
    """Форматы экспорта"""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    JSON = "json"
    CSV = "csv"

class ExportType(Enum):
    """Типы экспорта"""
    CHAT_HISTORY = "chat_history"
    DOCUMENT_ANALYSIS = "document_analysis"
    USER_STATISTICS = "user_statistics"
    SUBSCRIPTION_REPORT = "subscription_report"
    PAYMENT_REPORT = "payment_report"
    ANNOTATIONS_REPORT = "annotations_report"
    REFERRAL_REPORT = "referral_report"

@dataclass
class ExportRequest:
    """Запрос на экспорт"""
    id: str
    user_id: int
    export_type: ExportType
    export_format: ExportFormat
    filters: Dict[str, Any]
    created_at: datetime
    status: str = "pending"
    file_path: Optional[str] = None
    file_size: Optional[int] = None
    error_message: Optional[str] = None

@dataclass
class ExportData:
    """Данные для экспорта"""
    title: str
    data: List[Dict[str, Any]]
    metadata: Dict[str, Any]

class ExportManager:
    """Менеджер экспорта"""
    
    def __init__(self):
        self.export_requests: Dict[str, ExportRequest] = {}
        self.export_templates = self._initialize_templates()
    
    def _initialize_templates(self) -> Dict[ExportType, Dict[str, Any]]:
        """Инициализация шаблонов экспорта"""
        return {
            ExportType.CHAT_HISTORY: {
                "title": "История чата",
                "columns": ["Дата", "Вопрос", "Ответ", "Категория", "Тональность"],
                "filename_template": "chat_history_{date}.{format}"
            },
            ExportType.DOCUMENT_ANALYSIS: {
                "title": "Анализ документов",
                "columns": ["Документ", "Тип", "Статус", "Дата анализа", "Результат"],
                "filename_template": "document_analysis_{date}.{format}"
            },
            ExportType.USER_STATISTICS: {
                "title": "Статистика пользователя",
                "columns": ["Метрика", "Значение", "Дата", "Изменение"],
                "filename_template": "user_statistics_{date}.{format}"
            },
            ExportType.SUBSCRIPTION_REPORT: {
                "title": "Отчет по подпискам",
                "columns": ["Пользователь", "Тариф", "Дата начала", "Дата окончания", "Статус"],
                "filename_template": "subscription_report_{date}.{format}"
            },
            ExportType.PAYMENT_REPORT: {
                "title": "Отчет по платежам",
                "columns": ["Дата", "Сумма", "Тип", "Статус", "Описание"],
                "filename_template": "payment_report_{date}.{format}"
            },
            ExportType.ANNOTATIONS_REPORT: {
                "title": "Отчет по аннотациям",
                "columns": ["Документ", "Тип аннотации", "Содержимое", "Дата создания", "Автор"],
                "filename_template": "annotations_report_{date}.{format}"
            },
            ExportType.REFERRAL_REPORT: {
                "title": "Отчет по рефералам",
                "columns": ["Реферер", "Приведенный", "Тип", "Комиссия", "Дата", "Статус"],
                "filename_template": "referral_report_{date}.{format}"
            }
        }
    
    def create_export_request(
        self,
        user_id: int,
        export_type: ExportType,
        export_format: ExportFormat,
        filters: Optional[Dict[str, Any]] = None
    ) -> ExportRequest:
        """Создание запроса на экспорт"""
        try:
            request_id = f"export_{int(datetime.now().timestamp())}_{uuid.uuid4().hex[:8]}"
            
            request = ExportRequest(
                id=request_id,
                user_id=user_id,
                export_type=export_type,
                export_format=export_format,
                filters=filters or {},
                created_at=datetime.now(),
                status="pending"
            )
            
            self.export_requests[request_id] = request
            
            logger.info(f"Created export request {request_id} for user {user_id}")
            return request
            
        except Exception as e:
            logger.error(f"Export request creation error: {str(e)}")
            raise
    
    def process_export(self, request_id: str) -> bool:
        """Обработка запроса на экспорт"""
        try:
            request = self.export_requests.get(request_id)
            if not request:
                return False
            
            request.status = "processing"
            
            # Получаем данные для экспорта
            export_data = self._get_export_data(request)
            
            # Экспортируем в нужном формате
            file_path = self._export_to_format(request, export_data)
            
            if file_path:
                request.status = "completed"
                request.file_path = file_path
                request.file_size = self._get_file_size(file_path)
                logger.info(f"Export completed: {request_id}")
                return True
            else:
                request.status = "failed"
                request.error_message = "Failed to create export file"
                return False
                
        except Exception as e:
            logger.error(f"Export processing error: {str(e)}")
            if request_id in self.export_requests:
                self.export_requests[request_id].status = "failed"
                self.export_requests[request_id].error_message = str(e)
            return False
    
    def _get_export_data(self, request: ExportRequest) -> ExportData:
        """Получение данных для экспорта"""
        # В реальном приложении здесь будут запросы к базе данных
        # Пока используем тестовые данные
        
        template = self.export_templates.get(request.export_type, {})
        
        if request.export_type == ExportType.CHAT_HISTORY:
            data = [
                {
                    "Дата": "2024-01-15 10:30:00",
                    "Вопрос": "Как оформить трудовой договор?",
                    "Ответ": "Трудовой договор оформляется в письменной форме...",
                    "Категория": "Трудовое право",
                    "Тональность": "Нейтральная"
                },
                {
                    "Дата": "2024-01-15 11:15:00",
                    "Вопрос": "Какие документы нужны для регистрации ООО?",
                    "Ответ": "Для регистрации ООО необходимы следующие документы...",
                    "Категория": "Корпоративное право",
                    "Тональность": "Положительная"
                }
            ]
        elif request.export_type == ExportType.USER_STATISTICS:
            data = [
                {
                    "Метрика": "Всего сообщений",
                    "Значение": "150",
                    "Дата": "2024-01-15",
                    "Изменение": "+10"
                },
                {
                    "Метрика": "Активные дни",
                    "Значение": "25",
                    "Дата": "2024-01-15",
                    "Изменение": "+2"
                }
            ]
        else:
            data = []
        
        return ExportData(
            title=template.get("title", "Отчет"),
            data=data,
            metadata={
                "export_date": datetime.now().isoformat(),
                "user_id": request.user_id,
                "filters": request.filters
            }
        )
    
    def _export_to_format(self, request: ExportRequest, data: ExportData) -> Optional[str]:
        """Экспорт в нужном формате"""
        try:
            if request.export_format == ExportFormat.PDF:
                return self._export_to_pdf(request, data)
            elif request.export_format == ExportFormat.DOCX:
                return self._export_to_docx(request, data)
            elif request.export_format == ExportFormat.XLSX:
                return self._export_to_xlsx(request, data)
            elif request.export_format == ExportFormat.JSON:
                return self._export_to_json(request, data)
            elif request.export_format == ExportFormat.CSV:
                return self._export_to_csv(request, data)
            else:
                return None
                
        except Exception as e:
            logger.error(f"Export to {request.export_format.value} error: {str(e)}")
            return None
    
    def _export_to_pdf(self, request: ExportRequest, data: ExportData) -> str:
        """Экспорт в PDF"""
        from app.core.config import settings
        import os
        filename = os.path.join(settings.TEMP_DIR, f"export_{request.id}.pdf")
        os.makedirs(settings.TEMP_DIR, exist_ok=True)
        
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Заголовок
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Центрирование
        )
        story.append(Paragraph(data.title, title_style))
        story.append(Spacer(1, 12))
        
        # Метаданные
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey
        )
        story.append(Paragraph(f"Дата экспорта: {data.metadata['export_date']}", meta_style))
        story.append(Spacer(1, 20))
        
        # Таблица данных
        if data.data:
            table_data = []
            
            # Заголовки
            if data.data:
                headers = list(data.data[0].keys())
                table_data.append(headers)
                
                # Данные
                for row in data.data:
                    table_data.append([str(row.get(header, "")) for header in headers])
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
        
        doc.build(story)
        return filename
    
    def _export_to_docx(self, request: ExportRequest, data: ExportData) -> str:
        """Экспорт в DOCX"""
        from app.core.config import settings
        import os
        filename = os.path.join(settings.TEMP_DIR, f"export_{request.id}.docx")
        os.makedirs(settings.TEMP_DIR, exist_ok=True)
        
        doc = Document()
        
        # Заголовок
        doc.add_heading(data.title, 0)
        
        # Метаданные
        doc.add_paragraph(f"Дата экспорта: {data.metadata['export_date']}")
        doc.add_paragraph()
        
        # Таблица данных
        if data.data:
            table = doc.add_table(rows=1, cols=len(data.data[0].keys()))
            table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = list(data.data[0].keys())
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # Данные
            for row_data in data.data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(row_data.get(header, ""))
        
        doc.save(filename)
        return filename
    
    def _export_to_xlsx(self, request: ExportRequest, data: ExportData) -> str:
        """Экспорт в XLSX"""
        from app.core.config import settings
        import os
        filename = os.path.join(settings.TEMP_DIR, f"export_{request.id}.xlsx")
        os.makedirs(settings.TEMP_DIR, exist_ok=True)
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = data.title
        
        # Заголовок
        ws['A1'] = data.title
        ws['A1'].font = openpyxl.styles.Font(size=16, bold=True)
        
        # Метаданные
        ws['A3'] = f"Дата экспорта: {data.metadata['export_date']}"
        
        # Таблица данных
        if data.data:
            headers = list(data.data[0].keys())
            
            # Заголовки
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=5, column=col, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
            
            # Данные
            for row, row_data in enumerate(data.data, 6):
                for col, header in enumerate(headers, 1):
                    ws.cell(row=row, column=col, value=row_data.get(header, ""))
        
        wb.save(filename)
        return filename
    
    def _export_to_json(self, request: ExportRequest, data: ExportData) -> str:
        """Экспорт в JSON"""
        from app.core.config import settings
        import os
        filename = os.path.join(settings.TEMP_DIR, f"export_{request.id}.json")
        os.makedirs(settings.TEMP_DIR, exist_ok=True)
        
        export_data = {
            "title": data.title,
            "metadata": data.metadata,
            "data": data.data
        }
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)
        
        return filename
    
    def _export_to_csv(self, request: ExportRequest, data: ExportData) -> str:
        """Экспорт в CSV"""
        from app.core.config import settings
        import os
        filename = os.path.join(settings.TEMP_DIR, f"export_{request.id}.csv")
        os.makedirs(settings.TEMP_DIR, exist_ok=True)
        
        import csv
        
        with open(filename, 'w', newline='', encoding='utf-8') as f:
            if data.data:
                headers = list(data.data[0].keys())
                writer = csv.DictWriter(f, fieldnames=headers)
                writer.writeheader()
                writer.writerows(data.data)
        
        return filename
    
    def _get_file_size(self, file_path: str) -> int:
        """Получение размера файла"""
        try:
            import os
            return os.path.getsize(file_path)
        except Exception:
            return 0
    
    def get_export_request(self, request_id: str) -> Optional[ExportRequest]:
        """Получение запроса на экспорт"""
        return self.export_requests.get(request_id)
    
    def get_user_export_requests(self, user_id: int) -> List[ExportRequest]:
        """Получение запросов на экспорт пользователя"""
        user_requests = [req for req in self.export_requests.values() if req.user_id == user_id]
        user_requests.sort(key=lambda x: x.created_at, reverse=True)
        return user_requests
    
    def get_export_stats(self) -> Dict[str, Any]:
        """Получение статистики экспорта"""
        total_requests = len(self.export_requests)
        completed_requests = len([req for req in self.export_requests.values() if req.status == "completed"])
        failed_requests = len([req for req in self.export_requests.values() if req.status == "failed"])
        
        # Статистика по форматам
        format_stats = {}
        for format_type in ExportFormat:
            count = len([req for req in self.export_requests.values() if req.export_format == format_type])
            format_stats[format_type.value] = count
        
        # Статистика по типам
        type_stats = {}
        for export_type in ExportType:
            count = len([req for req in self.export_requests.values() if req.export_type == export_type])
            type_stats[export_type.value] = count
        
        return {
            "total_requests": total_requests,
            "completed_requests": completed_requests,
            "failed_requests": failed_requests,
            "success_rate": completed_requests / total_requests if total_requests > 0 else 0,
            "format_distribution": format_stats,
            "type_distribution": type_stats
        }

# Глобальный экземпляр менеджера экспорта
export_manager = ExportManager()

def get_export_manager() -> ExportManager:
    """Получение экземпляра менеджера экспорта"""
    return export_manager
